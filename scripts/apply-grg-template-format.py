#!/usr/bin/env python3
"""Add GRG template placeholders to a .docx (intro only + scan marker)."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

GRG_DATE = "{{GRG_DATE}}"
GRG_SONG_LIST = "{{GRG_SONG_LIST}}"
GRG_SCANS_BEGIN = "{{GRG_SCANS_BEGIN}}"
ROSTER_NAME_POSITION = "[Name | First-name Last Initial]: [Position]"

# Scan-section style exemplars. Each token paragraph is formatted exactly like
# its golden counterpart; the fill engine reads these styles AND the column
# layout of the section each token sits in, then inserts plaintext that adopts
# them. Sizes are points (golden: title 14, credits 9, body 12). The `columns`
# field defines the section layout: header (1 col) vs lyrics (2 cols). Edit
# these (sizes, colors, and the column counts) to restyle all scans.
STYLE_EXEMPLARS = [
    # token, size_pt, bold, color_rgb, highlight, font_name, columns
    ("{{STYLE_TITLE}}", 14, True, None, None, None, 1),
    ("{{STYLE_CREDIT}}", 9, False, None, None, "Helvetica Neue", 1),
    ("{{STYLE_BAR}}", 12, True, RGBColor(0xFF, 0x00, 0x00), None, None, 2),
    ("{{STYLE_LABEL}}", 12, True, None, WD_COLOR_INDEX.YELLOW, None, 2),
    ("{{STYLE_LYRIC}}", 12, False, None, None, None, 2),
]

MONTH_DATE = re.compile(
    r"(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d",
    re.I,
)


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = text


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def paragraph_has_page_break(paragraph) -> bool:
    for br in paragraph._element.xpath(".//w:br"):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def find_indices(doc: Document) -> dict[str, int]:
    date_idx = None
    song_header_idx = None
    first_key_idx = None
    last_key_idx = None
    page_break_idx = None

    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if date_idx is None and MONTH_DATE.search(t):
            date_idx = i
        if "song list" in t.lower():
            song_header_idx = i
        if t.lower().startswith("key of"):
            if first_key_idx is None:
                first_key_idx = i
            last_key_idx = i
        if page_break_idx is None and paragraph_has_page_break(p):
            page_break_idx = i

    if date_idx is None:
        raise ValueError("Could not find date paragraph (month name + day).")
    if first_key_idx is None or last_key_idx is None:
        raise ValueError('Could not find "Key of …" song list lines.')
    if page_break_idx is None:
        raise ValueError("Could not find page break after song list.")

    return {
        "date": date_idx,
        "song_header": song_header_idx,
        "first_key": first_key_idx,
        "last_key": last_key_idx,
        "page_break": page_break_idx,
    }


ROSTER_LINE_RE = re.compile(r"^\[Name\s*\|\s*[^\]]*\]:", re.I)
BAND_HEADER_RE = re.compile(r"^BAND\s*[\(:]", re.I)
CHOIR_HEADER_RE = re.compile(r"^CHOIR\s*[\(:]", re.I)
ALL_TEAM_HEADER_RE = re.compile(r"^ALL\s+TEAM", re.I)


def apply_roster_placeholders(doc: Document) -> None:
    """One [Name]: [Position] line per BAND and CHOIR section; remove SUN-specific position rows."""
    current_section: str | None = None
    section_first_roster_idx: dict[str, int] = {}
    roster_idxs_to_delete: list[int] = []

    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if BAND_HEADER_RE.match(t):
            current_section = "band"
            continue
        if CHOIR_HEADER_RE.match(t):
            current_section = "choir"
            continue
        if ALL_TEAM_HEADER_RE.match(t):
            current_section = "all_team"
            continue
        if "song list" in t.lower():
            current_section = None
            continue

        if current_section in ("band", "choir") and ROSTER_LINE_RE.match(t):
            if current_section not in section_first_roster_idx:
                section_first_roster_idx[current_section] = i
            else:
                roster_idxs_to_delete.append(i)

    for i in sorted(roster_idxs_to_delete, reverse=True):
        delete_paragraph(doc.paragraphs[i])

    for section, idx in section_first_roster_idx.items():
        set_paragraph_text(doc.paragraphs[idx], ROSTER_NAME_POSITION)


def apply_template_format(doc: Document) -> None:
    apply_roster_placeholders(doc)

    idx = find_indices(doc)

    set_paragraph_text(doc.paragraphs[idx["date"]], GRG_DATE)

    set_paragraph_text(doc.paragraphs[idx["first_key"]], GRG_SONG_LIST)
    for i in range(idx["last_key"], idx["first_key"], -1):
        if i != idx["first_key"]:
            delete_paragraph(doc.paragraphs[i])

    # Re-find page break after deletions
    page_break_idx = None
    for i, p in enumerate(doc.paragraphs):
        if paragraph_has_page_break(p):
            page_break_idx = i
            break
    if page_break_idx is None:
        raise ValueError("Page break missing after song list edit.")

    scans_marker_idx = page_break_idx + 1
    if scans_marker_idx < len(doc.paragraphs):
        set_paragraph_text(doc.paragraphs[scans_marker_idx], GRG_SCANS_BEGIN)
    else:
        doc.add_paragraph(GRG_SCANS_BEGIN)

    # Remove all content after marker (sample scans not needed in template)
    while len(doc.paragraphs) > scans_marker_idx + 1:
        delete_paragraph(doc.paragraphs[scans_marker_idx + 1])

    add_style_exemplars(doc)


def set_section_columns(section, num: int, space: int = 720) -> None:
    """Set the column count on a section's sectPr (w:cols)."""
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space))


def add_style_exemplars(doc: Document) -> None:
    """Append styled {{STYLE_*}} exemplar paragraphs grouped into column sections.

    The fill engine reads each token's run style AND the column layout of the
    section it sits in, then inserts plaintext that adopts both. Header tokens
    live in a single-column section; lyric tokens in a two-column section. Edit
    these (and their column counts) in the Drive template to restyle all scans.
    """
    # The trailing section (marker + first exemplars) must be single column;
    # the golden's body-final section is two-column and would otherwise bleed.
    set_section_columns(doc.sections[-1], 1)

    prev_columns = 1
    for token, size_pt, bold, color_rgb, highlight, font_name, columns in STYLE_EXEMPLARS:
        if columns != prev_columns:
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            set_section_columns(section, columns)
            prev_columns = columns

        paragraph = doc.add_paragraph()
        run = paragraph.add_run(token)
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        if font_name is not None:
            run.font.name = font_name
        if color_rgb is not None:
            run.font.color.rgb = color_rgb
        if highlight is not None:
            run.font.highlight_color = highlight


def main() -> int:
    src = Path(
        sys.argv[1]
        if len(sys.argv) > 1
        else "/Users/SBBWD/Downloads/Get Ready Guide (SUN).docx",
    )
    dst = Path(sys.argv[2]) if len(sys.argv) > 2 else src

    doc = Document(str(src))
    apply_template_format(doc)
    doc.save(str(dst))
    print(f"Saved template-formatted doc: {dst}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
